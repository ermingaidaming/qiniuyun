from __future__ import annotations

import io

from docx import Document
from docx.shared import Inches, Pt

from app.models.screenplay import Screenplay
from app.services.screenyaml import screenyaml_dumps


def export_yaml(screenplay: Screenplay) -> bytes:
    """Export screenplay as ScreenYAML format."""
    return screenyaml_dumps(screenplay).encode("utf-8")


def export_txt(screenplay: Screenplay) -> bytes:
    """Export screenplay as plain text."""
    lines: list[str] = []
    lines.append(screenplay.title)
    lines.append("=" * len(screenplay.title))
    lines.append("")

    for scene in screenplay.scenes:
        lines.append(f"场景 {scene.index}: {scene.setting}")
        lines.append("-" * 40)

        for elem in scene.elements:
            content = elem.content
            if elem.type == "character":
                lines.append(f"\n  [{content}]")
            elif elem.type == "dialogue":
                lines.append(f"    {elem.character or ''}: \"{content}\"")
            elif elem.type == "parenthetical":
                lines.append(f"    ({content})")
            elif elem.type == "action":
                lines.append(f"  {content}")

        lines.append("")

    return "\n".join(lines).encode("utf-8")


def export_docx(screenplay: Screenplay) -> bytes:
    """Export screenplay as a formatted Word document."""
    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = 1  # Center
    title_run = title_para.add_run(screenplay.title)
    title_run.bold = True
    title_run.font.size = Pt(18)

    doc.add_paragraph()  # Spacer

    for scene in screenplay.scenes:
        # Scene heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run(f"场景 {scene.index}: {scene.setting}")
        heading_run.bold = True
        heading_run.font.size = Pt(13)

        for elem in scene.elements:
            if elem.type == "action":
                para = doc.add_paragraph()
                para.paragraph_format.first_line_indent = Inches(0.3)
                para.add_run(elem.content).font.size = Pt(11)

            elif elem.type == "character":
                para = doc.add_paragraph()
                para.alignment = 1  # Center
                run = para.add_run(elem.content)
                run.bold = True
                run.font.size = Pt(12)

            elif elem.type == "dialogue":
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(1)
                para.paragraph_format.right_indent = Inches(1)
                para.add_run(elem.content).font.size = Pt(11)

            elif elem.type == "parenthetical":
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(1.5)
                run = para.add_run(f"({elem.content})")
                run.italic = True
                run.font.size = Pt(10)

        doc.add_paragraph()  # Scene separator

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
